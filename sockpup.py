import os
import requests
import getpass
import random
from docx.oxml import parse_xml
from docx import Document
from docx.shared import Inches
from io import BytesIO
from faker import Faker
from datetime import date

art = """
   _____            __      ____                         __     
  / ___/____  _____/ /__   / __ \__  ______  ____  ___  / /_    
  \__ \/ __ \/ ___/ //_/  / /_/ / / / / __ \/ __ \/ _ \/ __/    
 ___/ / /_/ / /__/ ,<    / ____/ /_/ / /_/ / /_/ /  __/ /_      
/____/\____/\___/_/|_|  /_/    \__,_/ .___/ .___/\___/\__/      
       ______                      /_/  _/_/                    
      / ____/__  ____  ___  _________ _/ /_____  _____          
     / / __/ _ \/ __ \/ _ \/ ___/ __ `/ __/ __ \/ ___/          
    / /_/ /  __/ / / /  __/ /  / /_/ / /_/ /_/ / /              
    \____/\___/_/ /_/\___/_/   \__,_/\__/\____/_/               
"""

fake = Faker()
user = getpass.getuser()
command = "/sockpup"
use = "Sock puppet generator"

def fetch_avatar():
    response = requests.get("https://thispersondoesnotexist.com/")
    response.raise_for_status()
    return BytesIO(response.content)

def create_sock_puppet():
    profile = {
        "name": fake.name(),
        "username": fake.user_name(),
        "email": fake.email(),
        "address": {
            "street": fake.street_name(),
            "city": fake.city(),
            "state": fake.state(),
            "zipcode": fake.zipcode(),
            "country": fake.country()
        },
        "phone": fake.phone_number(),
        "birthdate": fake.date_of_birth(tzinfo=None, minimum_age=18, maximum_age=65),
        "website": fake.url(),
        "place_of_birth": fake.city(),
        "height": f"{random.randint(150, 200)} cm", # In cm
        "weight": f"{random.randint(50, 100)} kg", # In kg
        "eye_color": fake.color_name(),
        "hobbies": ", ".join(fake.words(nb=3, unique=True)),
        "password": fake.password()
    }

    locale = fake.random_element(elements=('en_US', 'en_GB', 'en_AU', 'es_ES', 'de_DE', 'fr_FR'))
    fake_with_locale = Faker(locale)
    profile["phone"] = fake_with_locale.phone_number()
    profile["address"]["country"] = fake_with_locale.country()

    tech_companies = ['Google', 'Microsoft', 'Apple']
    finance_companies = ['Goldman Sachs', 'JP Morgan', 'Morgan Stanley']
    company_type = fake.random_element(elements=('tech', 'finance', 'none'))
    if company_type == 'tech':
        profile["company"] = fake.random_element(elements=tech_companies)
        profile["job_title"] = "Software Developer"  # Replace with your logic
    elif company_type == 'finance':
        profile["company"] = fake.random_element(elements=finance_companies)
        profile["job_title"] = "Financial Analyst"  # Replace with your logic
    else:
        profile["company"] = None
        profile["job_title"] = fake.job()

    age = (date.today() - profile['birthdate']).days // 365
    if age < 25:
        profile["job_title"] = "Junior " + profile["job_title"]
    elif age > 50:
        profile["job_title"] = "Senior " + profile["job_title"]

    return profile

def print_profile(profile):
    print(f"Name: {profile['name']}")
    print(f"Username: {profile['username']}")
    print(f"Email: {profile['email']}")
    print(f"Phone: {profile['phone']}")
    if profile["company"]:
        print(f"Company: {profile['company']}")
    print(f"Address: {profile['address']['street']}, {profile['address']['city']}, {profile['address']['state']}, {profile['address']['zipcode']}, {profile['address']['country']}")
    print(f"Birthdate: {profile['birthdate']}")
    print(f"Job Title: {profile['job_title']}")
    print(f"Website: {profile['website']}")
    print(f"Place of Birth: {profile['place_of_birth']}")
    print(f"Height: {profile['height']}")
    print(f"Weight: {profile['weight']}")
    print(f"Eye Color: {profile['eye_color']}")
    print(f"Hobbies: {profile['hobbies']}")
    print(f"Password: {profile['password']}")

def export_to_docx(profile, avatar_stream, save_path):
    doc = Document()
    doc.add_heading(profile['name'], level=0)

    table = doc.add_table(rows=1, cols=2)

    # Remove table outlines
    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:noBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />'))

    # Add profile picture to the left cell
    cell_1 = table.cell(0, 0)
    cell_1_paragraph = cell_1.paragraphs[0]
    cell_1_paragraph.add_run().add_picture(avatar_stream, width=Inches(2.0))

    # Add profile details to the right cell
    cell_2 = table.cell(0, 1)
    details = f"Name: {profile['name']}\nUsername: {profile['username']}\nEmail: {profile['email']}\nPhone: {profile['phone']}"
    if profile["company"]:
        details += f"\nCompany: {profile['company']}"
    cell_2.text = details

    doc.add_paragraph(f"\nAddress: {profile['address']['street']}, {profile['address']['city']}, {profile['address']['state']}, {profile['address']['zipcode']}, {profile['address']['country']}")
    doc.add_paragraph(f"Birthdate: {profile['birthdate']}")
    doc.add_paragraph(f"\nPlace of Birth: {profile['place_of_birth']}")
    doc.add_paragraph(f"Height: {profile['height']}")
    doc.add_paragraph(f"Weight: {profile['weight']}")
    doc.add_paragraph(f"Eye Color: {profile['eye_color']}")
    doc.add_paragraph(f"Job Title: {profile['job_title']}")
    doc.add_paragraph(f"Website: {profile['website']}")
    doc.add_paragraph(f"Password: {profile['password']}")
    doc.add_paragraph(f"Hobbies: {profile['hobbies']}")
    
    doc.save(save_path)

def run():
    print(art)
    print("Navi> Lets build you some new identification...")
    while True:
        profile = create_sock_puppet()
        avatar_stream = fetch_avatar()
        print_profile(profile)
        print("\nOptions:")
        print("[n] Generate a new profile")
        print("[s] Save this profile to a docx file")
        print("[e] Exit")
        
        choice = input("Choose an option: ")
        
        if choice == 'n':
            print()
            continue
        elif choice == 's':
            save_path = input(f"Navi> Enter the location and filename to save (ex: '/home/{user}/profile.docx'):\n")
            directory = os.path.dirname(save_path)
            if directory == '' or os.path.isdir(directory):
                export_to_docx(profile, avatar_stream, save_path)
                print(f"Navi> Profile saved to '{save_path}'")
                break
            else:
                print("Navi> Invalid directory. Please ensure the directory exists and try again.")
        elif choice == 'e':
            break
        else:
            print("Navi> Invalid option. Please choose again.")

